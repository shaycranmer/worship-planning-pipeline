#!/usr/bin/env python3
"""
Worship Plan Generator
Takes unordered input (songs, chunks, metadata) and assembles into proper template order.
Outputs formatted document ready for final editing.
"""

import os
from datetime import date
from typing import Dict, List, Optional
from worship_database import WorshipDatabase
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Template A: Thanksgiving for Baptism
TEMPLATE_A = {
    "name": "Thanksgiving for Baptism",
    "structure": [
        # GATHERING
        ("section_header", "GATHERING"),
        ("rubric", "The Holy Spirit calls us together as the people of God."),
        ("chunk", "welcome", "Welcome", "contextual"),
        ("chunk", "thanksgiving_baptism", "Thanksgiving for Baptism", "custom_with_intro"),
        ("song", "entrance", "Entrance Song"),  # Optional — Palm Sunday procession etc.
        ("song", "gathering", "Gathering Song"),
        ("chunk", "greeting", "Greeting", "standard"),
        ("chunk", "peace", "Peace", "standard"),
        ("chunk", "prayer_of_day", "Prayer of the Day", "custom_with_intro"),

        # WORD
        ("section_header", "WORD"),
        ("rubric", "God speaks to us in scripture reading, preaching, and song."),
        ("reading", "first_reading", "First Reading"),
        ("song", "gospel_acclamation", "Gospel Acclamation"),
        ("reading", "gospel", "Gospel"),
        ("chunk", "sermon", "Sermon", "contextual"),
        ("song", "hymn_of_day", "Hymn of the Day"),
        ("chunk", "creed", "Creed", "standard_or_options"),
        ("chunk", "prayers", "Prayers of Intercession", "custom"),
        ("section_header", "Announcements"),

        # MEAL
        ("section_header", "MEAL"),
        ("song", "offering", "Offering Song"),
        ("chunk", "offering_prayer", "Offering Prayer", "custom_with_intro"),
        ("chunk", "great_thanksgiving", "Great Thanksgiving / Dialogue", "standard"),
        ("chunk", "preface", "Preface", "standard"),
        ("song", "sanctus", "Holy, Holy, Holy (Sanctus)"),  # Optional — skipped if not provided, standard liturgical text used
        ("chunk", "thanksgiving_table", "Thanksgiving at the Table (Words of Institution)", "standard"),
        ("chunk", "lords_prayer", "Lord's Prayer", "template_specific"),
        ("chunk", "lamb_of_god", "Lamb of God (Agnus Dei)", "standard"),
        ("song", "communion", "Communion Song"),
        ("chunk", "prayer_after_communion", "Prayer after Communion", "custom_with_intro"),

        # SENDING
        ("section_header", "SENDING"),
        ("chunk", "blessing", "Blessing", "standard_or_custom"),
        ("song", "sending", "Sending Song"),
        ("chunk", "dismissal", "Dismissal", "standard_or_custom"),
    ]
}

# Template B: Confession and Kyrie
TEMPLATE_B = {
    "name": "Confession and Kyrie",
    "structure": [
        # GATHERING
        ("section_header", "GATHERING"),
        ("rubric", "The Holy Spirit calls us together as the people of God."),
        ("chunk", "welcome", "Welcome", "contextual"),
        ("song", "entrance", "Entrance Song"),  # Optional — Palm Sunday procession etc.
        ("song", "gathering", "Gathering Song"),
        ("chunk", "confession", "Confession and Forgiveness", "standard_or_custom"),
        ("chunk", "kyrie", "Kyrie", "standard"),
        ("chunk", "greeting", "Greeting", "standard"),
        ("chunk", "peace", "Peace", "standard"),
        ("chunk", "prayer_of_day", "Prayer of the Day", "custom_with_intro"),

        # WORD
        ("section_header", "WORD"),
        ("rubric", "God speaks to us in scripture reading, preaching, and song."),
        ("reading", "first_reading", "First Reading"),
        ("song", "gospel_acclamation", "Gospel Acclamation"),
        ("reading", "gospel", "Gospel"),
        ("chunk", "sermon", "Sermon", "contextual"),
        ("song", "hymn_of_day", "Hymn of the Day"),
        ("chunk", "creed", "Creed", "standard_or_options"),
        ("chunk", "prayers", "Prayers of Intercession", "custom"),
        ("section_header", "Announcements"),

        # MEAL
        ("section_header", "MEAL"),
        ("song", "offering", "Offering Song"),
        ("chunk", "offering_prayer", "Offering Prayer", "custom_with_intro"),
        ("chunk", "great_thanksgiving", "Great Thanksgiving / Dialogue", "standard"),
        ("chunk", "preface", "Preface", "standard"),
        ("song", "sanctus", "Holy, Holy, Holy (Sanctus)"),  # Optional — skipped if not provided, standard liturgical text used
        ("chunk", "thanksgiving_table", "Thanksgiving at the Table (Words of Institution)", "standard"),
        ("chunk", "lords_prayer", "Lord's Prayer", "template_specific"),
        ("chunk", "lamb_of_god", "Lamb of God (Agnus Dei)", "standard"),
        ("song", "communion", "Communion Song"),
        ("chunk", "prayer_after_communion", "Prayer after Communion", "custom_with_intro"),

        # SENDING
        ("section_header", "SENDING"),
        ("chunk", "blessing", "Blessing", "standard_or_custom"),
        ("song", "sending", "Sending Song"),
        ("chunk", "dismissal", "Dismissal", "standard_or_custom"),
    ]
}

# Standard Lutheran liturgical texts
STANDARD_TEXTS = {
    "greeting": """The grace of our Lord Jesus Christ, the love of God,
and the communion of the Holy Spirit be with you all.
**And also with you.**""",

    "peace": """The peace of Christ be with you always.
**And also with you.**""",

    "great_thanksgiving": """The Lord be with you.
**And also with you.**
Lift up your hearts.
**We lift them to the Lord.**
Let us give thanks to the Lord our God.
**It is right to give our thanks and praise.**""",

    "preface": """It is indeed right, our duty and our joy,
that we should at all times and in all places
give thanks and praise to you, almighty and merciful God,
through our Savior Jesus Christ;
who on this day overcame death and the grave,
and by his glorious resurrection opened to us the way of everlasting life.
And so, with all the choirs of angels,
with the church on earth and the hosts of heaven,
we praise your name and join their unending hymn:""",

    "holy": """Holy, holy, holy Lord,
God of pow'r and might,
heaven and earth are full of your glory, full of your glory.
Hosanna, hosanna, hosanna in the highest.
Blessed is he who comes in the name of the Lord.
Hosanna, hosanna, hosanna in the highest.""",

    "thanksgiving_table": """In the night in which he was betrayed,
our Lord Jesus took bread, and gave thanks;
broke it, and gave it to his disciples, saying:
Take and eat; this is my body, given for you.
Do this for the remembrance of me.

Again, after supper, he took the cup, gave thanks,
and gave it for all to drink, saying:
This cup is the new covenant in my blood,
shed for you and for all people for the forgiveness of sin.
Do this for the remembrance of me.""",

    "lords_prayer_traditional": """**Our Father, who art in heaven,
hallowed be thy name,
thy kingdom come,
thy will be done,
on earth as it is in heaven.
Give us this day our daily bread;
and forgive us our trespasses,
as we forgive those who trespass against us;
and lead us not into temptation,
but deliver us from evil.
For thine is the kingdom,
and the power, and the glory,
forever and ever. Amen.**""",

    "lords_prayer_modern": """**Our Father in heaven,
hallowed be your name,
your kingdom come,
your will be done,
on earth as in heaven.
Give us today our daily bread.
Forgive us our sins
as we forgive those who sin against us.
Save us from the time of trial
and deliver us from evil.
For the kingdom, the power,
and the glory are yours
now and for ever. Amen.**""",

    "lamb_of_god": """Lamb of God, you take away the sin of the world; have mercy on us. Lamb of God, you take away the sin of the world; have mercy on us. Lamb of God, you take away the sin of the world; grant us peace, grant us peace, Lamb of God.""",

    "dismissal": """Go in peace. Serve the Lord.
**Thanks be to God.**""",

    "thanksgiving_baptism": """Blessed be the holy Trinity, ☩ one God,
the fountain of living water,
the rock who gave us birth,
our light and our salvation.
**Amen.**

Joined to Christ in the waters of baptism,
we are clothed with God's mercy and forgiveness.
Let us give thanks for the gift of baptism.

We give you thanks, O God,
for in the beginning your Spirit moved over the waters
and by your Word you created the world,
calling forth life in which you took delight.

Through the waters of the flood you delivered Noah and his family.
Through the sea you led your people Israel from slavery into freedom.
At the river your Son was baptized by John and anointed with the Holy Spirit.
By water and your Word you claim us as daughters and sons,
making us heirs of your promise and servants of all.

We praise you for the gift of water that sustains life,
and above all we praise you for the gift of new life in Jesus Christ.
Shower us with your Spirit,
and renew our lives with your forgiveness, grace, and love.

To you be given honor and praise
through Jesus Christ our Lord
in the unity of the Holy Spirit, now and forever.
**Amen.**""",

    "confession": """In the name of the Father,
and of the ☩ Son,
and of the Holy Spirit.
**Amen.**

God of all mercy and consolation, come to the help of your people, turning us from our sin to live for you alone. Give us the power of your Holy Spirit that we may confess our sin, receive your forgiveness, and grow into the fullness of Jesus Christ, our Savior and Lord.
Amen.

Let us confess our sin in the presence of God and of one another.

[Silence is kept for reflection]

**Gracious God,
have mercy on us. We confess that we have turned from you and given ourselves into the power of sin. We are truly sorry and humbly repent. In your compassion forgive us our sins, known and unknown, things we have done and things we have failed to do. Turn us again to you, and uphold us by your Spirit, so that we may live and serve you in newness of life through Jesus Christ, our Savior and Lord.
Amen.**

God, who is rich in mercy, loved us even when we were dead in sin, and made us alive together with Christ. By grace you have been saved. In the name of ☩ Jesus Christ, your sins are forgiven. Almighty God strengthen you with power through the Holy Spirit, that Christ may live in your hearts through faith.
**Amen.**""",

    "kyrie": """Kyrie eleison, on our world and on our way,
Kyrie eleison, ev'ry day.
For peace in the world, for the health of the church, for the unity of all;
for this holy house, for all who worship and praise,
let us pray to the Lord, let us pray to the Lord
Kyrie eleison, on our world and on our way,
Kyrie eleison, ev'ry day.
That we may live out your impassioned response to the hungry and the poor;
that we may live out truth and justice and grace,
let us pray to the Lord, let us pray to the Lord
Kyrie eleison, on our world and on our way,
Kyrie eleison, ev'ry day.
For peace in our hearts, for peace in our homes, for friends and family;
for life and for love, for our work and our play,
let us pray to the Lord, let us pray to the Lord
Kyrie eleison, on our world and on our way,
Kyrie eleison, ev'ry day.
For your Spirit to guide; that you center our lives in the water and the Word;
that you nourish our souls with your body and blood,
let us pray to the Lord, let us pray to the Lord
Kyrie eleison, on our world and on our way,
Kyrie eleison, ev'ry day.""",

    "creed": """**I believe in God, the Father almighty,
creator of heaven and earth.
I believe in Jesus Christ, God's only Son, our Lord,
who was conceived by the Holy Spirit,
born of the virgin Mary,
suffered under Pontius Pilate,
was crucified, died, and was buried;
he descended to the dead.
On the third day he rose again;
he ascended into heaven,
he is seated at the right hand of the Father,
and he will come to judge the living and the dead.
I believe in the Holy Spirit,
the holy catholic church,
the communion of saints,
the forgiveness of sins,
the resurrection of the body,
and the life everlasting. Amen.**""",
}


class WorshipPlanGenerator:
    """Generates formatted worship plans from unordered input."""

    def __init__(self, database_path: str = "worship.db"):
        self.db = WorshipDatabase(database_path)
        self.db_path = database_path
        self.templates = {
            'A': TEMPLATE_A,
            'B': TEMPLATE_B
        }

    def close(self):
        """Close database connection."""
        self.db.close()

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    def generate_plan(self,
                     template: str,
                     service_date: date,
                     season: str = "",
                     theme: str = "",
                     readings: Dict[str, str] = None,
                     songs: Dict[str, str] = None,
                     custom_chunks: Dict[str, str] = None) -> str:
        """
        Generate worship plan from unordered input.

        Args:
            template: 'A' or 'B'
            service_date: Date of worship service
            season: Liturgical season (Advent, Lent, Easter, etc.)
            theme: Worship theme
            readings: Dict mapping reading slots to text (e.g., {'first_reading': 'Joel 2:12-13'})
            songs: Dict mapping song slots to song titles (e.g., {'gathering': 'Be Thou My Vision'})
            custom_chunks: Dict mapping chunk IDs to custom text

        Returns:
            Formatted worship plan as string
        """
        if readings is None:
            readings = {}
        if songs is None:
            songs = {}
        if custom_chunks is None:
            custom_chunks = {}

        template_data = self.templates.get(template.upper())
        if not template_data:
            raise ValueError(f"Unknown template: {template}")

        # Build the plan
        lines = []
        lines.append("=" * 80)
        lines.append(f"WORSHIP PLAN - {service_date.strftime('%B %d, %Y')}")
        if season:
            lines.append(f"Season: {season}")
        if theme:
            lines.append(f"Theme: {theme}")
        lines.append("=" * 80)
        lines.append("")

        # Process template structure
        for item in template_data['structure']:
            item_type = item[0]

            if item_type == "section_header":
                section_name = item[1]
                lines.append("")
                lines.append(section_name)
                lines.append("-" * len(section_name))
                lines.append("")

            elif item_type == "rubric":
                # Add rubric text
                rubric_text = item[1]
                lines.append(rubric_text)
                lines.append("")

            elif item_type == "chunk":
                chunk_id = item[1]
                chunk_label = item[2]
                chunk_type = item[3]

                lines.append(f"{chunk_label}:")
                lines.append("")

                # Check if custom text provided
                if chunk_id in custom_chunks:
                    lines.append(self._indent(custom_chunks[chunk_id]))
                # Check if standard text available
                elif chunk_type == "standard" and chunk_id in STANDARD_TEXTS:
                    lines.append(self._indent(STANDARD_TEXTS[chunk_id]))
                elif chunk_type == "standard_or_custom" and chunk_id in STANDARD_TEXTS:
                    lines.append(self._indent(STANDARD_TEXTS[chunk_id]))
                    lines.append(self._indent("[Or use custom text]"))
                else:
                    lines.append(self._indent(f"[{chunk_label} - to be written]"))

                lines.append("")

            elif item_type == "song":
                slot_id = item[1]
                slot_label = item[2]

                # Entrance is truly optional — skip entirely if not provided
                if slot_id == "entrance" and slot_id not in songs:
                    continue

                # Sanctus: if no song provided, render standard liturgical text
                if slot_id == "sanctus" and slot_id not in songs:
                    lines.append(f"{slot_label}:")
                    lines.append("")
                    lines.append(self._indent(STANDARD_TEXTS["holy"]))
                    lines.append("")
                    continue

                lines.append(f"{slot_label}:")
                lines.append("")

                if slot_id in songs:
                    song_title = songs[slot_id]
                    # Look up song in database
                    song_data = self._find_song(song_title)

                    if song_data:
                        # Format song info
                        if song_data['type'] == 'traditional' and song_data['hymnal_number']:
                            lines.append(self._indent(f"{song_data['title']} ({song_data['hymnal_number']})"))
                        else:
                            lines.append(self._indent(f"{song_data['title']} - {song_data['artist_source']}"))

                        # Add lyrics if available (placeholder for now)
                        lines.append(self._indent("[Lyrics to be added]"))
                    else:
                        lines.append(self._indent(f"{song_title} [Not found in database]"))
                else:
                    lines.append(self._indent(f"[{slot_label} - to be selected]"))

                lines.append("")

            elif item_type == "reading":
                reading_id = item[1]
                reading_label = item[2]

                lines.append(f"{reading_label}:")
                lines.append("")

                if reading_id in readings:
                    reading_text = readings[reading_id]

                    # Extract book name from reading citation
                    book_name = reading_text.split()[0] if reading_text else "Scripture"

                    # Add reader cue
                    if reading_id == "gospel":
                        lines.append(self._indent(f"The holy gospel according to {book_name}."))
                        lines.append(self._indent("Congregation responds: Glory to you, O Lord."))
                    else:
                        lines.append(self._indent(f"A reading from {book_name}."))

                    lines.append("")
                    lines.append(self._indent(reading_text))
                    lines.append("")

                    # Add response
                    if reading_id == "gospel":
                        lines.append(self._indent("The gospel of the Lord.\nCongregation: Praise to you, O Christ."))
                    else:
                        lines.append(self._indent("The word of the Lord.\nCongregation: Thanks be to God."))
                else:
                    lines.append(self._indent(f"[{reading_label} - to be added]"))

                lines.append("")

        return "\n".join(lines)

    def _find_song(self, title: str) -> Optional[Dict]:
        """Find song in database by title."""
        songs = self.db.find_songs_by_theme(title)
        # Try exact match first
        for song in songs:
            if song['title'].lower() == title.lower():
                return song
        # Return first partial match if no exact
        return songs[0] if songs else None

    def _indent(self, text: str, spaces: int = 4) -> str:
        """Indent text block."""
        indent = " " * spaces
        return "\n".join(indent + line for line in text.split("\n"))

    def _get_song_lyrics(self, song_id: int) -> Optional[str]:
        """Fetch lyrics for a song directly from database by ID."""
        import sqlite3
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT lyrics FROM songs WHERE id = ?', (song_id,))
        row = cursor.fetchone()
        conn.close()
        if row and row[0] and row[0].strip():
            return row[0]
        return None

    def _clean_lyrics(self, lyrics: str, title: str) -> str:
        """Clean lyric-bank text for display: decode RTF escapes, drop a duplicate
        leading title line, and strip trailing copyright/attribution lines."""
        import re

        # Decode \'XX RTF/codepage escapes (e.g. \'e9 -> é, \'a9 -> ©)
        def _dec(m):
            try:
                return bytes([int(m.group(1), 16)]).decode("cp1252")
            except Exception:
                return m.group(0)
        lyrics = re.sub(r"\\'([0-9a-fA-F]{2})", _dec, lyrics)

        lines = lyrics.split("\n")

        # Drop a leading line that just repeats the song title (avoids double title)
        def _norm(s):
            return re.sub(r"[^a-z0-9 ]", "", s.lower()).strip()
        if lines and title:
            first = lines[0].strip()
            nl, nt = _norm(first), _norm(title)
            # A title header carries no sentence punctuation; a real first lyric line
            # usually does ("Let the vineyards be fruitful, Lord,"). Only treat a prefix
            # match as a title if it looks like a bare header.
            looks_like_header = not re.search(r"[,.;:!?]", first)
            title_match = bool(nl) and (nl == nt or (looks_like_header and (nl.startswith(nt) or nt.startswith(nl))))
            # ...but if that line recurs later, it's a sung lyric, not a header
            # ("Halle Halle Hallelujah" is both the title and the first sung line).
            recurs = any(_norm(l) == nl for l in lines[1:])
            if title_match and not recurs:
                lines = lines[1:]

        # Strip trailing attribution / copyright lines
        attrib = re.compile(r"^(text[:\s]|music[:\s]|words[:\s]|tune[:\s]|©|\(c\)|copyright)", re.I)
        rights = re.compile(r"(all rights reserved|admin\.|publications|augsburg|giamusic|"
                            r"ocp |used by permission|lutheran book of worship)", re.I)
        while lines and (not lines[-1].strip() or attrib.search(lines[-1].strip()) or rights.search(lines[-1])):
            lines.pop()

        return "\n".join(lines).strip()

    def _add_formatted_paragraph(self, doc, text: str, indent_inches: float = 0.5,
                                 bold_responses: bool = True):
        """Add text with proper paragraph breaks and bold formatting for **text**.

        bold_responses auto-bolds liturgical responses (Amen, Thanks be to God). Turn
        it OFF for scripture, which may contain those words non-liturgically (e.g.
        Jeremiah 28:6, "Amen! May the Lord do so.")."""
        import re

        if bold_responses:
            # strip any existing wrap first, then re-apply so it's safe
            for _pat in (r'Amen\.?', r'Thanks be to God\.?'):
                text = re.sub(r'\*\*\s*(' + _pat + r')\s*\*\*', r'\1', text, flags=re.IGNORECASE)
                text = re.sub(r'\b(' + _pat + r')', r'**\1**', text, flags=re.IGNORECASE)

        # Split into paragraphs on blank lines
        paragraphs = [p for p in text.split('\n\n') if p.strip()]

        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent_inches)

            # Split by ** markers for bold text, handling line breaks within
            parts = para_text.split('**')
            for i, part in enumerate(parts):
                is_bold = (i % 2 == 1)
                lines = part.split('\n')
                for j, line in enumerate(lines):
                    if j > 0:
                        p.add_run('\n')
                    if line:
                        run = p.add_run(line)
                        if is_bold:
                            run.bold = True

    def _add_prayers_paragraph(self, doc, text: str, indent_inches: float = 0.5):
        """Add prayers text with paragraph spacing and auto-bold congregational responses."""
        import re

        # Normalize first — strip any existing ** around response phrases to avoid double-wrapping
        text = re.sub(r'\*\*(hear our prayer\.?)\*\*', r'\1', text, flags=re.IGNORECASE)
        text = re.sub(r'\*\*(Lord, in your mercy\.?)\*\*', r'\1', text, flags=re.IGNORECASE)

        # Now apply bold consistently
        text = re.sub(r'(hear our prayer\.?)', r'**\1**', text, flags=re.IGNORECASE)
        text = re.sub(r'(Lord, in your mercy\.?)', r'**\1**', text, flags=re.IGNORECASE)

        # Now render with paragraph breaks
        self._add_formatted_paragraph(doc, text, indent_inches)

    def generate_docx(self,
                     template: str,
                     service_date: date,
                     season: str = "",
                     theme: str = "",
                     readings: Dict[str, str] = None,
                     songs: Dict[str, str] = None,
                     custom_chunks: Dict[str, str] = None,
                     filename: str = None) -> str:
        """
        Generate worship plan as formatted Word document.

        Returns path to generated file.
        """
        if readings is None:
            readings = {}
        if songs is None:
            songs = {}
        if custom_chunks is None:
            custom_chunks = {}

        if filename is None:
            filename = f"worship_plan_{service_date.isoformat()}.docx"

        # Output to Worship Plans folder if it exists and filename has no path
        if not os.path.dirname(filename):
            plans_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Worship Plans")
            if os.path.isdir(plans_dir):
                filename = os.path.join(plans_dir, filename)

        template_data = self.templates.get(template.upper())
        if not template_data:
            raise ValueError(f"Unknown template: {template}")

        # Create document
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_heading(f"Worship Plan - {service_date.strftime('%B %d, %Y')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if season or theme:
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if season:
                meta.add_run(f"Season: {season}").italic = True
            if season and theme:
                meta.add_run(" | ")
            if theme:
                meta.add_run(f"Theme: {theme}").italic = True

        doc.add_paragraph()  # Spacing

        # Process template structure
        for item in template_data['structure']:
            item_type = item[0]

            if item_type == "section_header":
                section_name = item[1]
                heading = doc.add_heading(section_name, level=1)
                heading.runs[0].font.size = Pt(14)

            elif item_type == "rubric":
                # Add rubric text (liturgical instruction)
                rubric_text = item[1]
                p = doc.add_paragraph(rubric_text)
                p.runs[0].italic = True
                doc.add_paragraph()  # spacing

            elif item_type == "chunk":
                chunk_id = item[1]
                chunk_label = item[2]
                chunk_type = item[3]

                # Chunk label
                p = doc.add_paragraph()
                p.add_run(f"{chunk_label}").bold = True

                # Chunk content
                # Contextual chunks (welcome, sermon) - no placeholder text
                if chunk_type == "contextual":
                    if chunk_id in custom_chunks:
                        self._add_formatted_paragraph(doc, custom_chunks[chunk_id])
                    # Otherwise just leave the label, no placeholder

                # Custom with intro (prayers with "Let us pray")
                elif chunk_type == "custom_with_intro":
                    import re as _re
                    p = doc.add_paragraph("Let us pray.")
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.runs[0].italic = True
                    if chunk_id in custom_chunks:
                        # Strip a leading "Let us pray" so it isn't printed twice
                        body = _re.sub(r'^\s*Let us pray[.,]?\s*\n?', '',
                                       custom_chunks[chunk_id], flags=_re.IGNORECASE)
                        self._add_formatted_paragraph(doc, body)
                    elif chunk_id in STANDARD_TEXTS:
                        # Fall back to standard ELW text (e.g. thanksgiving_baptism)
                        self._add_formatted_paragraph(doc, STANDARD_TEXTS[chunk_id])

                # Template-specific (Lord's Prayer switches based on template)
                elif chunk_type == "template_specific":
                    if chunk_id == "lords_prayer":
                        # Add intro rubric
                        p = doc.add_paragraph("Gathered into one by the Holy Spirit, let us pray as Jesus taught us.")
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.runs[0].italic = True
                        # Choose prayer based on template
                        if template.upper() == 'A':
                            self._add_formatted_paragraph(doc, STANDARD_TEXTS["lords_prayer_traditional"])
                        else:
                            self._add_formatted_paragraph(doc, STANDARD_TEXTS["lords_prayer_modern"])

                # Standard text only
                elif chunk_type == "standard" and chunk_id in STANDARD_TEXTS:
                    self._add_formatted_paragraph(doc, STANDARD_TEXTS[chunk_id])

                # Standard or custom - use custom if provided, otherwise standard
                elif chunk_type == "standard_or_custom":
                    if chunk_id in custom_chunks:
                        self._add_formatted_paragraph(doc, custom_chunks[chunk_id])
                    elif chunk_id in STANDARD_TEXTS:
                        self._add_formatted_paragraph(doc, STANDARD_TEXTS[chunk_id])

                # Standard or options (like Creed) - use custom if provided, otherwise standard
                elif chunk_type == "standard_or_options":
                    if chunk_id in custom_chunks:
                        self._add_formatted_paragraph(doc, custom_chunks[chunk_id])
                    elif chunk_id in STANDARD_TEXTS:
                        self._add_formatted_paragraph(doc, STANDARD_TEXTS[chunk_id])

                # Prayers - special rendering with paragraph spacing and auto-bold responses
                elif chunk_id == "prayers" and chunk_id in custom_chunks:
                    self._add_prayers_paragraph(doc, custom_chunks[chunk_id])

                # Custom text only
                elif chunk_id in custom_chunks:
                    self._add_formatted_paragraph(doc, custom_chunks[chunk_id])

                else:
                    p = doc.add_paragraph(f"[{chunk_label} - to be written]")
                    p.paragraph_format.left_indent = Inches(0.5)

            elif item_type == "song":
                slot_id = item[1]
                slot_label = item[2]

                # Entrance is truly optional — skip entirely if not provided
                if slot_id == "entrance" and slot_id not in songs:
                    continue

                # Sanctus: if no song provided, render standard liturgical text
                if slot_id == "sanctus" and slot_id not in songs:
                    p = doc.add_paragraph()
                    p.add_run(f"{slot_label}:").bold = True
                    self._add_formatted_paragraph(doc, STANDARD_TEXTS["holy"])
                    continue

                # Song label
                p = doc.add_paragraph()
                p.add_run(f"{slot_label}:").bold = True

                if slot_id in songs:
                    song_title = songs[slot_id]
                    song_data = self._find_song(song_title)

                    if song_data:
                        # Song title
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        if song_data['hymnal_number']:
                            p.add_run(f"{song_data['title']} ({song_data['hymnal_number']})").italic = True
                        else:
                            p.add_run(f"{song_data['title']} - {song_data['artist_source']}").italic = True

                        # Pull lyrics from database if available, otherwise placeholder
                        lyrics = self._get_song_lyrics(song_data['id'])
                        if lyrics:
                            lyrics = self._clean_lyrics(lyrics, song_data['title'])
                            p = doc.add_paragraph(lyrics)
                            p.paragraph_format.left_indent = Inches(0.75)
                        else:
                            p = doc.add_paragraph("[Lyrics to be added]")
                            p.paragraph_format.left_indent = Inches(0.75)
                    else:
                        p = doc.add_paragraph(f"{song_title} [Not found in database]")
                        p.paragraph_format.left_indent = Inches(0.5)
                else:
                    p = doc.add_paragraph(f"[{slot_label} - to be selected]")
                    p.paragraph_format.left_indent = Inches(0.5)

            elif item_type == "reading":
                reading_id = item[1]
                reading_label = item[2]

                if reading_id in readings:
                    reading_text = readings[reading_id]

                    # Parse citation and body text.
                    # Full scripture uses \n\n as separator (fetch_readings format).
                    # Short citations use " - " (worship_input.json format).
                    if '\n\n' in reading_text:
                        parts = reading_text.split('\n\n', 1)
                        citation = parts[0].strip()
                        description = parts[1].strip()
                    elif ' - ' in reading_text:
                        citation, description = reading_text.split(' - ', 1)
                        citation = citation.strip()
                        description = description.strip()
                    else:
                        citation = reading_text.strip()
                        description = None

                    # Extract book name from citation (e.g., "Genesis 2:15-17")
                    book_name = citation.split()[0] if citation else "Scripture"

                    # Reading label with citation on same line (bold)
                    p = doc.add_paragraph()
                    p.add_run(f"{reading_label}: {citation}").bold = True

                    # Reader cue
                    if reading_id == "gospel":
                        p = doc.add_paragraph(f"The holy gospel according to {book_name}.")
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.runs[0].italic = True

                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.add_run("Glory to you, O Lord.").bold = True
                    else:
                        p = doc.add_paragraph(f"A reading from {book_name}.")
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.runs[0].italic = True

                    # Description / full scripture text if present
                    if description:
                        self._add_formatted_paragraph(doc, description, indent_inches=0.5,
                                                      bold_responses=False)

                    # Response
                    if reading_id == "gospel":
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.add_run("The gospel of the Lord.").italic = True
                        p.add_run("\n")
                        p.add_run("Praise to you, O Christ.").bold = True
                    else:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.add_run("The word of the Lord.").italic = True
                        p.add_run("\n")
                        p.add_run("Thanks be to God.").bold = True
                else:
                    p = doc.add_paragraph(f"[{reading_label} - to be added]")
                    p.paragraph_format.left_indent = Inches(0.5)

        # Save document
        doc.save(filename)
        return filename

    def suggest_songs(self, theme_keywords: List[str], slots: List[str]) -> Dict[str, List[Dict]]:
        """
        Suggest songs for each slot based on theme and rotation.

        Args:
            theme_keywords: List of theme words to search for
            slots: List of song slot IDs (e.g., ['gathering', 'hymn_of_day', 'offering'])

        Returns:
            Dict mapping slot IDs to list of suggested songs
        """
        suggestions = {}

        for slot in slots:
            # Different logic per slot type
            if slot == 'gospel_acclamation':
                # Usually "Halle, Halle, Hallelujah"
                suggestions[slot] = self.db.find_songs_by_theme("Halle")[:1]

            elif slot == 'hymn_of_day':
                # Traditional hymn matching theme
                all_matches = []
                for keyword in theme_keywords:
                    matches = self.db.find_songs_by_theme(keyword)
                    all_matches.extend([s for s in matches if s['type'] == 'traditional'])

                # Remove duplicates and get fresh songs
                seen = set()
                unique = []
                for song in all_matches:
                    if song['id'] not in seen:
                        seen.add(song['id'])
                        unique.append(song)

                suggestions[slot] = unique[:3]

            elif slot in ['entrance', 'gathering', 'offering', 'sending']:
                # Contemporary songs matching theme
                all_matches = []
                for keyword in theme_keywords:
                    matches = self.db.find_songs_by_theme(keyword)
                    all_matches.extend([s for s in matches if s['type'] == 'contemporary'])

                # Remove duplicates
                seen = set()
                unique = []
                for song in all_matches:
                    if song['id'] not in seen:
                        seen.add(song['id'])
                        unique.append(song)

                suggestions[slot] = unique[:3]

            elif slot == 'communion':
                # Multiple communion songs, mix of types
                all_matches = []
                for keyword in theme_keywords + ['communion', 'table', 'meal']:
                    matches = self.db.find_songs_by_theme(keyword)
                    all_matches.extend(matches)

                seen = set()
                unique = []
                for song in all_matches:
                    if song['id'] not in seen:
                        seen.add(song['id'])
                        unique.append(song)

                suggestions[slot] = unique[:5]

        return suggestions


def interactive_plan():
    """Interactive worship plan builder."""
    print("\n" + "="*60)
    print("WORSHIP PLAN GENERATOR")
    print("="*60)

    gen = WorshipPlanGenerator()

    try:
        # Get basic info
        print("\n--- BASIC INFORMATION ---")
        template = input("Template (A=Thanksgiving for Baptism, B=Confession): ").upper()

        date_str = input("Service date (YYYY-MM-DD) or Enter for next Sunday: ")
        if date_str:
            service_date = date.fromisoformat(date_str)
        else:
            # Default to next Sunday (simplified)
            service_date = date.today()

        season = input("Season (Advent, Lent, Easter, etc.): ")
        theme = input("Theme: ")

        # Get readings
        print("\n--- READINGS ---")
        first_reading = input("First Reading (e.g., Joel 2:12-13): ")
        gospel = input("Gospel Reading: ")

        readings = {}
        if first_reading:
            readings['first_reading'] = first_reading
        if gospel:
            readings['gospel'] = gospel

        # Suggest songs
        print("\n--- SONG SUGGESTIONS ---")
        print("Enter theme keywords (comma-separated):")
        keywords_input = input("Keywords: ")
        keywords = [k.strip() for k in keywords_input.split(',')] if keywords_input else []

        if keywords:
            slots = ['gathering', 'hymn_of_day', 'offering', 'communion', 'sending']
            suggestions = gen.suggest_songs(keywords, slots)

            print("\nSuggested songs:")
            for slot, song_list in suggestions.items():
                print(f"\n{slot.upper()}:")
                for i, song in enumerate(song_list, 1):
                    if song['type'] == 'traditional':
                        print(f"  {i}. {song['title']} ({song['hymnal_number']})")
                    else:
                        print(f"  {i}. {song['title']} - {song['artist_source']}")

        # Select songs
        print("\n--- SELECT SONGS ---")
        songs = {}
        for slot in ['gathering', 'hymn_of_day', 'offering', 'sending']:
            song_title = input(f"{slot.capitalize()} song: ")
            if song_title:
                songs[slot] = song_title

        # Communion can have multiple
        communion_input = input("Communion song(s) (comma-separated): ")
        if communion_input:
            songs['communion'] = communion_input

        # Generate plan
        print("\n--- GENERATING PLAN ---")

        # Generate Word document
        docx_file = gen.generate_docx(
            template=template,
            service_date=service_date,
            season=season,
            theme=theme,
            readings=readings,
            songs=songs
        )

        print(f"\n✓ Worship plan saved to: {docx_file}")
        print("\nYou can now open this file in Microsoft Word or Google Docs.")
        print("Fill in any [to be written] sections and add lyrics as needed.")

    finally:
        gen.close()


if __name__ == "__main__":
    interactive_plan()
